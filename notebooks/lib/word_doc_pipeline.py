#!/usr/bin/env python3
"""
Reusable helpers for building and post-processing Word documents in the tax notes pipeline.

Features
--------
* Load a new document from an optional Word template (``.docx`` or ``.dotx``).
* Apply East Asian font metadata to runs that rely on CJK typefaces.
* Insert a Table of Contents field that can later be refreshed by Microsoft Word.
* Update Word fields on macOS through an ``osascript`` automation routine.
* Provide a small CLI utility for inserting a TOC and refreshing fields on an
  existing ``.docx`` file.

Example
-------
>>> from pathlib import Path
>>> from notebooks.lib.word_doc_pipeline import (
...     load_document, apply_run_font, insert_table_of_contents, save_document
... )
>>> doc = load_document(Path("template.dotx"))
>>> heading = doc.add_heading("示範文件", level=1)
>>> for run in heading.runs:
...     apply_run_font(run)
>>> insert_table_of_contents(doc, title="目錄")
>>> save_document(doc, Path("output.docx"), auto_update=True)
"""
from __future__ import annotations

import argparse
import shutil
import subprocess
import textwrap
from pathlib import Path
from typing import Iterable, Optional
import tempfile
import zipfile

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DEFAULT_EAST_ASIA_FONT: Optional[str] = None

__all__ = [
    "DEFAULT_EAST_ASIA_FONT",
    "apply_run_font",
    "load_document",
    "clear_placeholder_content",
    "insert_table_of_contents",
    "update_docx_fields_with_word",
    "save_document",
]


# --------------------------------------------------------------------------- #
# Low-level helpers
# --------------------------------------------------------------------------- #

def apply_run_font(run, east_asia_font: Optional[str] = DEFAULT_EAST_ASIA_FONT) -> None:
    """
    Apply East Asian font metadata to a ``Run`` instance.

    Parameters
    ----------
    run:
        A python-docx ``Run`` object.
    east_asia_font:
        The font family name for CJK text. Defaults to ``標楷體``.
    """
    if east_asia_font is None:
        return

    run.font.name = east_asia_font
    r = run._element  # noqa: SLF001 - python-docx exposes the underlying element
    r_pr = getattr(r, "rPr", None)
    if r_pr is None or len(r_pr) == 0:
        r_pr = r.get_or_add_rPr()
    r_fonts = getattr(r_pr, "rFonts", None)
    if r_fonts is None or len(r_fonts) == 0:
        r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:ascii"), east_asia_font)
    r_fonts.set(qn("w:hAnsi"), east_asia_font)


def remove_paragraph(paragraph) -> None:
    """Remove a paragraph from the document while keeping the rest intact."""
    p = paragraph._element  # noqa: SLF001
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def clear_placeholder_content(doc: Document) -> None:
    """
    Delete the initial empty paragraph that python-docx adds when creating a
    document from a template. This keeps subsequent insertions clean.
    """
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text.strip():
        remove_paragraph(doc.paragraphs[0])


def _iter_paragraph_runs(document: Document) -> Iterable:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            yield run


# --------------------------------------------------------------------------- #
# Document creation helpers
# --------------------------------------------------------------------------- #

def load_document(template_path: Optional[Path] = None, *, east_asia_font: Optional[str] = DEFAULT_EAST_ASIA_FONT) -> Document:
    """
    Create a ``Document`` instance from a template if provided, otherwise start
    from a blank Word file.
    """
    temp_copy: Optional[Path] = None
    if template_path:
        template_path = Path(template_path)
        if template_path.suffix.lower() == ".dotx":
            temp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            temp.close()
            temp_copy = Path(temp.name)
            with zipfile.ZipFile(template_path) as src, zipfile.ZipFile(temp_copy, "w") as dst:
                for item in src.infolist():
                    data = src.read(item.filename)
                    if item.filename == "[Content_Types].xml":
                        data = data.replace(
                            b"application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                            b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                        )
                    dst.writestr(item, data)
            doc = Document(temp_copy)
            doc._temp_template_path = temp_copy  # type: ignore[attr-defined]
        else:
            doc = Document(template_path)
    else:
        doc = Document()
    clear_placeholder_content(doc)
    if east_asia_font:
        for run in _iter_paragraph_runs(doc):
            apply_run_font(run, east_asia_font=east_asia_font)
    return doc


def insert_table_of_contents(
    doc: Document,
    title: Optional[str] = "目錄",
    level_range: str = "1-3",
    heading_level: int = 1,
    *,
    east_asia_font: Optional[str] = DEFAULT_EAST_ASIA_FONT,
    page_break_after: bool = True,
) -> None:
    """
    Insert a Table of Contents field into the supplied document.

    Parameters mirror the configuration used inside the notebooks.
    """
    if title:
        heading = doc.add_heading(title, level=heading_level)
        for run in heading.runs:
            apply_run_font(run, east_asia_font=east_asia_font)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), f'TOC \\o "{level_range}" \\h \\z \\u')
    run._r.append(field)  # noqa: SLF001
    apply_run_font(run, east_asia_font=east_asia_font)

    if page_break_after:
        paragraph.add_run().add_break(WD_BREAK.PAGE)


def save_document(doc: Document, output_path: Path, *, auto_update: bool = False, update_kwargs: Optional[dict] = None) -> Path:
    """
    Persist the document and optionally trigger the Word field update routine.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    temp_template = getattr(doc, "_temp_template_path", None)
    if temp_template:
        Path(temp_template).unlink(missing_ok=True)
    if auto_update:
        kwargs = update_kwargs or {}
        update_docx_fields_with_word(output_path, **kwargs)
    return output_path


# --------------------------------------------------------------------------- #
# Word automation
# --------------------------------------------------------------------------- #

def update_docx_fields_with_word(docx_path: Path, enabled: bool = True) -> None:
    """
    Use macOS ``osascript`` automation to force Microsoft Word to refresh fields
    such as a Table of Contents or figure captions.
    """
    if not enabled:
        return

    if shutil.which("osascript") is None:
        print("⚠️ 找不到 osascript，略過 Word 欄位更新。")
        return

    if not docx_path.exists():
        raise FileNotFoundError(f"找不到檔案：{docx_path}")

    docx_posix = docx_path.as_posix().replace('"', '\\"')
    script = textwrap.dedent(
        f"""
        tell application "Microsoft Word"
            activate
            set docPath to "{docx_posix}"
            set theDoc to open file name docPath
            try
                tell theDoc
                    set fieldCount to count of fields
                    if fieldCount > 0 then
                        repeat with idx from 1 to fieldCount
                            update field (field idx)
                        end repeat
                    end if
                    save
                end tell
                close theDoc saving yes
            on error errMsg number errNum
                try
                    close theDoc saving no
                end try
                error errMsg number errNum
            end try
        end tell
        """
    )

    result = subprocess.run(["osascript", "-e", script], capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError("AppleScript 更新欄位失敗：" + (result.stderr or result.stdout))
    print("已透過 Word 更新欄位。")


# --------------------------------------------------------------------------- #
# CLI utility
# --------------------------------------------------------------------------- #

def _cli_insert_toc(args: argparse.Namespace) -> None:
    docx_path = Path(args.docx).resolve()
    if not docx_path.exists():
        raise FileNotFoundError(f"找不到 docx：{docx_path}")

    doc = Document(docx_path)
    clear_placeholder_content(doc)

    if args.insert_toc:
        insert_table_of_contents(
            doc,
            title=args.toc_title,
            level_range=args.toc_range,
            heading_level=args.toc_heading_level,
            east_asia_font=args.east_asia_font,
            page_break_after=args.page_break_after_toc,
        )

    save_document(doc, docx_path, auto_update=args.auto_update)
    print(f"處理完成：{docx_path}")


def build_cli_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Apply table of contents and refresh Word fields for an existing .docx file.",
    )
    parser.add_argument("docx", help="欲處理的 Word 檔案路徑")
    parser.add_argument("--no-insert-toc", dest="insert_toc", action="store_false", help="僅刷新欄位，不插入目錄")
    parser.add_argument("--toc-title", default="目錄", help="目錄標題（預設：目錄）")
    parser.add_argument("--toc-range", default="1-3", help="目錄層級範圍（預設：1-3）")
    parser.add_argument("--toc-heading-level", type=int, default=1, help="目錄標題階層（預設：1）")
    parser.add_argument("--no-page-break-after-toc", dest="page_break_after_toc", action="store_false", help="目錄後不自動換頁")
    parser.add_argument("--east-asia-font", default=DEFAULT_EAST_ASIA_FONT, help='中文字體（留空表示沿用模板設定）')
    parser.add_argument("--no-auto-update", dest="auto_update", action="store_false", help="不自動透過 Word 更新欄位")
    parser.set_defaults(insert_toc=True, page_break_after_toc=True, auto_update=True)
    return parser


def main(argv: Optional[list[str]] = None) -> None:
    parser = build_cli_parser()
    args = parser.parse_args(argv)
    _cli_insert_toc(args)


if __name__ == "__main__":  # pragma: no cover - CLI entry point
    main()
